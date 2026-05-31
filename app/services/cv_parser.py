import os
import re
import json
import logging
import time
from concurrent.futures import ThreadPoolExecutor, as_completed

import pdfplumber

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import win32com.client
except ImportError:
    win32com = None

from dotenv import load_dotenv
from openai import OpenAI
from app.core.database import SessionLocal
from app.models.candidate_model import Candidate
from app.models.cv_analysis_model import CvAnalysis

from app.services.cloudinary_service import upload_cv_to_cloudinary, check_cv_exists_on_cloudinary, stream_cv_from_cloudinary

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

UPLOAD_FOLDER = "app/uploads"


def sync_from_cloudinary():
    """
    Descarca din Cloudinary toate fisierele care nu exista local in app/uploads/.
    Apelat automat inainte de analiza pentru a asigura ca fisierele sunt disponibile.
    Returneaza numarul de fisiere sincronizate.
    """
    import cloudinary.api
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

    try:
        results = cloudinary.api.resources(
            resource_type="raw",
            max_results=500
        )
        remote_files = results.get("resources", [])
    except Exception as e:
        logger.warning(f"Nu s-au putut lista fisierele din Cloudinary: {e}")
        return 0

    synced = 0
    for resource in remote_files:
        filename = resource.get("public_id", "")
        if not filename:
            continue

        ext = os.path.splitext(filename)[1].lower()
        if ext not in (".pdf", ".docx", ".doc"):
            continue

        local_path = os.path.join(UPLOAD_FOLDER, filename)

        if os.path.exists(local_path):
            continue

        buffer, mime = stream_cv_from_cloudinary(filename)
        if buffer:
            with open(local_path, "wb") as f:
                f.write(buffer.read())
            synced += 1
            logger.info(f"Sincronizat din Cloudinary: {filename}")

    if synced > 0:
        logger.info(f"Sincronizare Cloudinary completa: {synced} fisiere descarcate local")
    return synced

TEMP_FOLDER = "/tmp/nexas_hr"
os.makedirs(TEMP_FOLDER, exist_ok=True)

JOB_PROFILES_PATH = "config/job_profiles.json"
client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=os.getenv("OPENROUTER_API_KEY"))

# ─── CACHE FUNCTIONS ──────────────────────────────────────────────────────────

def _normalize_job_key(job_title):
    return normalize_text(job_title).strip()


def get_cached_analysis(cv_file, job_title):
    key = _normalize_job_key(job_title)
    db = SessionLocal()
    try:
        row = db.query(CvAnalysis).filter(
            CvAnalysis.cv_file == cv_file,
            CvAnalysis.job_title_normalized == key
        ).first()
        if row:
            return json.loads(row.result_json)
        return None
    except Exception as e:
        logger.warning(f"Cache read error pentru {cv_file}: {e}")
        return None
    finally:
        db.close()


def save_analysis_to_cache(cv_file, job_title, data):
    key = _normalize_job_key(job_title)
    db = SessionLocal()
    try:
        existing = db.query(CvAnalysis).filter(
            CvAnalysis.cv_file == cv_file,
            CvAnalysis.job_title_normalized == key
        ).first()

        result_json = json.dumps(data, ensure_ascii=False)

        if existing:
            existing.result_json = result_json
        else:
            db.add(CvAnalysis(
                cv_file=cv_file,
                job_title_normalized=key,
                result_json=result_json
            ))

        db.commit()
    except Exception as e:
        logger.warning(f"Cache write error pentru {cv_file}: {e}")
        db.rollback()
    finally:
        db.close()


# ─── TEXT EXTRACTION ──────────────────────────────────────────────────────────

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        logger.info(f"Text extras din {pdf_path}: {len(text)} caractere")
    except Exception as e:
        logger.error(f"Eroare la citirea PDF {pdf_path}: {e}")
    return text


def extract_text_from_docx(docx_path):
    text = ""
    if Document is None:
        logger.error("Lipseste python-docx.")
        return text
    try:
        document = Document(docx_path)
        for paragraph in document.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        logger.info(f"Text extras din {docx_path}: {len(text)} caractere")
    except Exception as e:
        logger.error(f"Eroare la citirea DOCX {docx_path}: {e}")
    return text


def extract_text_from_doc(doc_path):
    text = ""
    if win32com is None:
        logger.error("Lipseste pywin32.")
        return text
    word = None
    document = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        document = word.Documents.Open(os.path.abspath(doc_path))
        text = document.Content.Text or ""
        logger.info(f"Text extras din {doc_path}: {len(text)} caractere")
    except Exception as e:
        logger.error(f"Eroare la citirea DOC {doc_path}: {e}")
    finally:
        try:
            if document:
                document.Close(False)
        except Exception:
            pass
        try:
            if word:
                word.Quit()
        except Exception:
            pass
    return text


def extract_text_from_file(file_path):
    file_lower = file_path.lower()
    if file_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    if file_lower.endswith(".docx"):
        return extract_text_from_docx(file_path)
    if file_lower.endswith(".doc"):
        return extract_text_from_doc(file_path)
    return ""


# ─── HELPERS ──────────────────────────────────────────────────────────────────

def normalize_text(value):
    value = str(value or "").lower()
    for a, b in [("ș","s"),("ş","s"),("ă","a"),("â","a"),("î","i"),("ț","t"),("ţ","t"),("."," "),("-"," "),("_"," "),("/"," ")]:
        value = value.replace(a, b)
    return " ".join(value.split())


def load_job_profiles():
    if not os.path.exists(JOB_PROFILES_PATH):
        logger.warning(f"Job profiles file not found at {JOB_PROFILES_PATH}")
        return []
    try:
        with open(JOB_PROFILES_PATH, "r", encoding="utf-8") as f:
            return json.load(f).get("profiles", [])
    except Exception as e:
        logger.error(f"Eroare job_profiles.json: {e}")
        return []


def match_job_profile(target_job):
    profiles = load_job_profiles()
    job = normalize_text(target_job)
    best, best_score = None, 0
    for p in profiles:
        candidates = [p.get("job_title",""), p.get("job_id","")] + p.get("aliases", []) + p.get("equivalent_roles", [])
        score = 0
        for item in candidates:
            item = normalize_text(item)
            if not item:
                continue
            if item == job:
                score += 100
            elif item in job or job in item:
                score += 70
            else:
                iw, jw = set(item.split()), set(job.split())
                if iw:
                    score += int((len(iw & jw) / len(iw)) * 45)
        if score > best_score:
            best, best_score = p, score
    return best if best_score >= 25 else None


def build_fallback_profile(target_job):
    job_lower = normalize_text(target_job)

    operational_kw = [
        "operator", "muncitor", "lucrator", "manipulant", "depozitar", "stivuitorist",
        "curier", "casier", "ospatar", "vanzator", "agent vanzari", "curatenie",
        "ingrijitor", "paznic", "agent paza", "spalator", "sorter", "ambalator",
        "picker", "packer", "necalificat", "zilier", "ajutor",
    ]
    skilled_trade_kw = [
        "electrician", "sudor", "lacatus", "mecanic auto", "mecanic", "instalator",
        "zugrav", "zugrav vopsitor", "tamplar", "zidar", "faiantar", "montator",
        "frigotehnist", "automatist", "mentenanta", "tehnic",
    ]
    senior_kw = [
        "coordonator", "manager", "director", "analist", "inginer", "contabil",
        "economist", "auditor", "consultant", "specialist", "expert", "sef",
        "responsabil hr", "hr", "jurist", "avocat", "financiar", "controller",
        "programator", "developer", "it", "arhitect", "medic", "farmacist",
        # Roluri profesionale/administrative care cer studii superioare
        "consilier", "inspector", "referent", "judecator", "procuror", "notar",
        "diplomat", "ofiter", "parlamentar", "prefect", "primar",
        "sef birou", "sef serviciu", "functionar public",
    ]

    is_operational = any(kw in job_lower for kw in operational_kw)
    is_skilled_trade = any(kw in job_lower for kw in skilled_trade_kw)
    is_senior = any(kw in job_lower for kw in senior_kw)

    if is_operational:
        return {
            "job_title": target_job, "domain": "Operational", "level": "Entry",
            "requires_higher_education": False,
            "must_have": ["seriozitate", "punctualitate", "disponibilitate program"],
            "nice_to_have": ["experienta in domeniu similar", "permis auto"],
            "reject_if_missing": [],
            "overqualified_risk": ["licenta", "master", "doctorat", "inginer", "economist", "jurist", "studii superioare"],
            "red_flags": [],
            "max_score_rules": [
                "REGULA CRITICA: Daca candidatul are studii superioare (licenta/master/doctorat), final_score MAXIM 35 — supracalificat, nu va ramane.",
                "Daca candidatul nu are nicio experienta de munca, final_score maxim 58.",
                "Candidatul cu experienta in domenii similare (depozit/logistica/munca fizica/vanzari) poate lua 60-90.",
            ],
        }
    elif is_skilled_trade:
        return {
            "job_title": target_job, "domain": "Meserie calificata", "level": "Skilled",
            "requires_higher_education": False,
            "must_have": ["calificare sau experienta practica in meserie"],
            "nice_to_have": ["certificate autorizatii", "experienta recenta"],
            "reject_if_missing": ["experienta practica sau calificare in meserie"],
            "overqualified_risk": ["master", "doctorat"],
            "red_flags": ["fara experienta practica", "doar studii teoretice"],
            "max_score_rules": [
                "Fara experienta sau calificare practica in meserie, final_score maxim 38.",
                "Candidat cu master/doctorat fara experienta practica, final_score maxim 35.",
            ],
        }
    elif is_senior:
        return {
            "job_title": target_job, "domain": "Calificat/Senior/Profesional", "level": "Senior",
            "requires_higher_education": True,
            "must_have": ["studii superioare relevante pentru domeniu", "experienta minima in domeniu"],
            "nice_to_have": ["certificari", "limbi straine", "experienta de conducere"],
            "reject_if_missing": ["studii superioare in domeniu relevant", "experienta in domeniu"],
            "overqualified_risk": [],
            "red_flags": ["fara studii superioare", "fara experienta relevanta in domeniu", "background complet diferit"],
            "max_score_rules": [
                "REGULA CRITICA: Fara studii superioare in domeniu DIRECT relevant, final_score MAXIM 35.",
                "REGULA CRITICA: Fara NICIO experienta in domeniu sau domenii conexe, final_score MAXIM 38.",
                "REGULA CRITICA: Daca background-ul candidatului e complet diferit de domeniu (ex: vanzator/stivuitorist pentru rol juridic/parlamentar/financiar), final_score MAXIM 32.",
                "Experienta partiala sau tangentiala in domeniu: scor 39-55.",
                "Candidat cu studii relevante dar fara experienta: scor 40-55.",
            ],
        }
    else:
        return {
            "job_title": target_job, "domain": "General/Profesional", "level": "Middle",
            "requires_higher_education": False,
            "must_have": ["experienta in domeniu relevant sau conexe"],
            "nice_to_have": ["studii superioare", "certificari relevante"],
            "reject_if_missing": [],
            "overqualified_risk": [],
            "red_flags": ["lipsa oricarei experiente in domeniu sau domenii conexe"],
            "max_score_rules": [
                "Daca candidatul nu are NICIO experienta in domeniu sau domenii conexe, final_score MAXIM 38.",
                "Daca experienta este tangentiala (transferabilitate slaba), final_score maxim 48.",
                "Daca experienta este partial relevanta, scor 49-65.",
            ],
        }


def detect_required_license(target_job, profile=None):
    text = normalize_text(target_job)
    if profile:
        text += " " + normalize_text(" ".join(profile.get("must_have", []) + profile.get("reject_if_missing", [])))
    if "permis ce" in text or "categoria ce" in text or "cat ce" in text or "tir" in text:
        return "CE"
    if "permis b" in text or "categoria b" in text or "cat b" in text or "sofer b" in text:
        return "B"
    if "permis c" in text or "categoria c" in text or "cat c" in text or "sofer c" in text:
        return "C"
    if "permis d" in text or "categoria d" in text or "cat d" in text:
        return "D"
    return ""


def build_prompt(text, target_job, job_requirements=None):
    profile = match_job_profile(target_job) or build_fallback_profile(target_job)
    required_license = detect_required_license(target_job, profile)

    level_map = {"Entry": "ENTRY", "Skilled": "JUNIOR/MIDDLE", "Senior": "SENIOR", "Middle": "MIDDLE"}
    job_level = level_map.get(profile.get("level", "Middle"), "MIDDLE")

    must_have = profile.get("must_have", [])
    reject_if_missing = profile.get("reject_if_missing", [])

    job_req_block = ""
    if job_requirements and job_requirements.strip():
        job_req_block = (
            "CERINȚE OBLIGATORII COMPLETATE DE HR (PRIORITATE MAXIMĂ — citește înainte de orice altceva):\n"
            + job_requirements.strip() + "\n\n"
        )

    cerinte_block = ""
    if required_license:
        cerinte_block += f"⚠️  PERMIS {required_license} OBLIGATORIU. Dacă lipsește din CV → SCOR MAX 30 + REJECT.\n"
    if reject_if_missing:
        cerinte_block += "OBLIGATORII (lipsă = REJECT): " + ", ".join(reject_if_missing) + "\n"
    if must_have:
        cerinte_block += "MUST HAVE: " + ", ".join(must_have) + "\n"
    if cerinte_block:
        cerinte_block = "CERINȚE DETECTATE DIN PROFIL POST:\n" + cerinte_block + "\n"

    schema = """{
  "candidate_name": "Exact din CV sau Necunoscut",
  "email": "",
  "phone": "",
  "companies": [],
  "positions_held": [],
  "current_position": "Ultimul titlu din CV",
  "recommended_role_for_candidate": "Jobul natural al candidatului bazat STRICT pe CV - IGNORA postul vizat - scrie INTOTDEAUNA ceva concret",

  "job_level_detected": "ENTRY|JUNIOR|MIDDLE|SENIOR|EXECUTIVE",
  "cv_quality_score": "CLEAR|DECENT|VAGUE|POOR",

  "experience_score": 0,
  "education_score": 0,
  "soft_skills_score": 0,
  "penalties_total": 0,
  "bonuses_total": 0,
  "scoring_breakdown": "exp X/40 + edu Y/25 + soft Z/20 + bonus B/5 - penalitati W = total",

  "final_score": 0,
  "recommendation": "STRONG_YES|YES|MAYBE|NO|REJECT",
  "confidence_level": "high|medium|low",

  "overqualification_risk": "LOW|MEDIUM|HIGH",
  "instability_risk": "LOW|MEDIUM|HIGH",
  "level_mismatch": "LOW|MEDIUM|HIGH",

  "years_experience": 0,
  "years_relevant_domain": 0,

  "strengths": [],
  "gaps": [],
  "red_flags": [],
  "transferable_skills": [],

  "decision_reason": "2-3 propozitii clare cu citate din CV care justifica scorul",
  "missing_requirements": [],
  "must_verify_by_phone": [],

  "recommended_next_action": "CALL_NOW|CALL_LATER|KEEP_FOR_OTHER_ROLE|REJECT",
  "priority": "HIGH|MEDIUM|LOW",

  "summary": "Rezumat HR curat, max 3 propozitii, fara jargon",
  "manager_summary": "Max 3 randuri pentru hiring manager cu decizia si motivul",
  "phone_call_script": "Text gata de folosit la telefon, in romana, natural, cu 3-4 intrebari cheie",
  "interview_questions": [],

  "retention_probability": 0,
  "growth_potential": "LOW|MEDIUM|HIGH",
  "better_role_match": "Alt rol mai potrivit sau gol",
  "reject_reason_internal": "Motiv intern doar daca NO sau REJECT",

  "candidate_cluster": "SUPER-CLUSTERUL candidatului bazat pe rolul sau NATURAL (recommended_role_for_candidate). Alege EXACT UNUL: blue_collar | constructii_tehnic | tech_ing | medical | business | support | educatie. GHID: blue_collar=HoReCa/Transport/Depozit/Retail/Paza/Agricultura(orice munca operationala-manuala); constructii_tehnic=Constructii/Instalatii/Sudura/AutoService; tech_ing=IT/Software/Inginerie/Proiectant; medical=Medic/Asistent/Farmacist/Kinetoterapeut/Psiholog; business=Contabil/Economist/Jurist/HR/Manager/SalesB2B/Financiar/Imobiliare; support=Admin/Secretariat/CallCenter/Marketing/ONG; educatie=Profesor/Trainer/Coach/Formator.",

  "job_cluster": "SUPER-CLUSTERUL POSTULUI VIZAT (job_title). Acelasi ghid. Exemple corecte: gestionar depozit=blue_collar, sef depozit=blue_collar, sofer TIR=blue_collar, bucatar=blue_collar, casier=blue_collar, paznic=blue_collar, contabil=business, jurist=business, specialist HR=business, manager=business, asistent medical=medical, medic=medical, farmacist=medical, programator=tech_ing, inginer=tech_ing, electrician constructii=constructii_tehnic, zugrav=constructii_tehnic, mecanic auto=constructii_tehnic, profesor=educatie, trainer=educatie, agent call center=support, marketing specialist=support."
}"""

    return f"""EȘTI UN SISTEM AUTOMAT DE EVALUARE HR CU CALIBRARE PIAȚA ROMÂNIEI.
Nu ești recruiter uman. Ești motor de scoring obiectiv care evaluează strict pe DATE din CV, nu pe potențial sau presupuneri.
Scorul final trebuie să fie DEFENSIBIL: "De ce scor 47?" → se explică prin citate exacte din CV.

═══════════════════════════════════════════════════════════
POSTUL VIZAT: {target_job}
NIVELUL DETECTAT: {job_level}
{job_req_block}{cerinte_block}═══════════════════════════════════════════════════════════
REGULA DE AUR: DOVEZI, NU SPECULAȚII
═══════════════════════════════════════════════════════════

ÎNAINTE DE ORICE SCORING, extrage EXACT din CV:
  - Nume, perioadă, angajator, titlu, durată în luni/ani
  - Dacă nu-i menționat clar = NU PRESUPUNE

DOVEZI VALIDE (acordă puncte):
  ✓ "5 ani analist financiar, 2019-2024"
  ✓ "Excel avansat — pivot tables, VBA"
  ✓ "Manager 8 oameni, creștere 40% revenue în 2023"
  ✓ "Permis C/CE din 2015, activ"

DOVEZI INVALIDE (0 puncte):
  ✗ "Sunt comunicativ și orientat spre rezultate"
  ✗ "Lucru bine în echipă"
  ✗ "Experiență în domeniu" (fără detalii = VAGUITATE)
  ✗ "Office avansat" fără specificație

DACĂ CV E VAGU (perioade neprecizate, angajatori fără nume, titluri vagi):
  → -8 puncte automat + cv_quality_score: VAGUE/POOR + must_verify obligatoriu + scor MAX 50

═══════════════════════════════════════════════════════════
FAZA 1: KNOCKOUT RULES (aplică ÎNAINTE de calculul de puncte)
═══════════════════════════════════════════════════════════

JOB OPERAȚIONAL/ENTRY (operator, depozitar, vânzător, paznic, casier, stivuitorist, curier, manipulant):
  ├─ Candidat cu licență/master/doctorat = SUPRACALIFICAT → SCOR MAX 32 + REJECT
  └─ Background total diferit fără nicio experiență similară → SCOR MAX 38

JOB CALIFICAT/SENIOR (analist, inginer, contabil, consilier, inspector, referent, HR, IT, medic, avocat, notar, judecator):
  ├─ Fără studii superioare în domeniu direct → SCOR MAX 30
  ├─ Fără nicio experiență în domeniu sau conexe → SCOR MAX 35
  ├─ Background complet diferit (ex: vânzător→Consilier Parlamentar, șofer→Inspector Fiscal, depozitar→Contabil) → SCOR MAX 25 + REJECT
  └─ Experiență < 1 an în domeniu vizat → SCOR MAX 45

JOB MESERIE CALIFICATĂ (electrician, sudor, mecanic, șofer TIR, instalator):
  ├─ Fără atestat/certificare legală obligatorie → SCOR MAX 15 + REJECT
  └─ Certificare expirată sau inactivă > 3 ani → SCOR MAX 25

MATRICE COMPATIBILITATE DOMENIU — 3 NIVELURI (aplică ÎNAINTE de orice calcul):

  NIVEL 0 — INCOMPATIBIL (scor MAX 28 + REJECT):
    Blue Collar (HoReCa/Transport/Depozit/Retail/Pază/Agricultură) → IT, Inginerie, Medical, Financiar, Juridic, HR, Educație
    IT/Inginerie → Medical, Blue Collar operațional (șofer, casier, depozit, HoReCa, pază)
    Medical → Tot ce nu e Medical sau Educație (parțial)
    Financiar/Juridic/HR → Blue Collar operațional, Construcții, Medical
    Educație → Blue Collar operațional, Construcții

  NIVEL 1 — UPSKILL (scor MAX 45, recomandare CONSIDER):
    Blue Collar → Construcții/Automotive (meserie conexă)
    Blue Collar → Support (Call Center, Admin, Marketing) — aptitudini transferabile
    Construcții → IT/Inginerie (dacă CV arată studii superioare tehnice)
    Construcții → Support (supraveghere, coordonare)
    IT/Inginerie → Construcții, Support (admin tehnic)
    IT/Inginerie → Educație (trainer tehnic)
    Medical → Educație (formator medical)
    Support → Blue Collar, Construcții, IT/Inginerie (reconversie)
    Educație → IT/Inginerie (trainer), Medical (educator medical)

  NIVEL 2 — COMPATIBIL (scoring normal):
    Același cluster sau super-cluster compatibil
    IT/Inginerie ↔ Financiar/Juridic/HR/Management
    Business (Financiar/Juridic/HR/Sales/Management) ↔ Support (Admin/Marketing/Call Center)
    Business ↔ Educație
    Support ↔ Educație

  EXCEPȚIE RECONVERSIE: Dacă CV-ul demonstrează reconversie activă documentată
  (cursuri relevante + experiență min 6 luni în domeniu nou) → ignoră regula și evaluează reconversia.

CERINȚĂ HARD LIPSĂ (oricare job):
  └─ Orice cerință din lista OBLIGATORII de mai sus lipsește → SCOR MAX 35 + REJECT
     NU scrie "merită o șansă" când cerința fundamentală lipsește.

═══════════════════════════════════════════════════════════
FAZA 2: SCORING EXPERIENȚĂ (0-40 PUNCTE)
═══════════════════════════════════════════════════════════

Ponderare temporală: ultimii 3 ani = 100% | 3-5 ani = 70% | >5 ani = 40%

Experiență ZERO în domeniu (ENTRY):
  Cu studii relevante: 20 | Fără studii relevante: 12 | 6-12 luni experiență: 30

JUNIOR (1-3 ani în domeniu direct):
  1 an: 32 | 2 ani: 36 | 3 ani: 40

MIDDLE (3-7 ani direct):
  3-4 ani: 37 | 5-6 ani: 39 | 7 ani: 40

SENIOR (7+ ani, cu rezultate concrete dovedite):
  7-10 ani stabili: 40 | 10+ ani + leadership: 40 | 10+ ani job-hopper: 30

PENALITĂȚI EXPERIENȚĂ:
  Pauză >6 luni neexplicată: -5 | Job <3 luni fără motiv: -3 per job
  3+ schimbări industrie în 5 ani: -5 | CV vag (perioade imprecise): -8

═══════════════════════════════════════════════════════════
FAZA 3: SCORING EDUCAȚIE & CERTIFICĂRI (0-25 PUNCTE)
═══════════════════════════════════════════════════════════

REGULA EDUCAȚIE (aplică în ordine):
  1. Postul cere licență → candidatul are licență în domeniu relevant: 25/25 (PUNCTAJ COMPLET)
  2. Postul cere licență → candidatul are master relevant: 25/25 (echivalent sau superior)
  3. Postul cere master → candidatul are master relevant: 25/25
  4. Postul cere master → candidatul are DOAR licență: 15/25 (penalizare pentru lipsă master explicit cerut)
  5. Postul NU menționează explicit studii → candidatul are licență relevantă: 15/25 (bonus)
  6. Studii parțial relevante (domeniu conex, nu direct): 10/25
  7. Lipsă completă studii superioare când postul le cere explicit: 0/25

NU penaliza licența dacă postul NU cere explicit master sau doctorat.

Certificări (se adaugă la scorul de educație, maxim 25 total):
  Certificare obligatorie prezentă: +8 | Certificare obligatorie LIPSĂ: -15
  Certificări relevante în domeniu (AWS, SAP, Excel avansat, CISA, CFA, etc.): +3 fiecare
  Cursuri recente relevante (sub 2 ani): +2

Limbi (acordă puncte DOAR dacă postul le cere și sunt menționate explicit):
  Nivel avansat dovedit: 5 | Nivel mediu menționat: 3 | Nemenționate: 0 — NU presupune engleză!

═══════════════════════════════════════════════════════════
FAZA 4: SOFT SKILLS & ATITUDINE (0-20 PUNCTE)
REGULA GENERALĂ: Dovezi concrete = puncte depline. Afirmații vagi = 0.
REGULA CV ROMÂNESC: CV-urile românești descriu adesea responsabilități, nu KPI-uri cu cifre.
  Dacă nu există KPI-uri cuantificate DAR există responsabilități clare + stabilitate angajator,
  acordă max 15/20 (nu 0). Aplică această regulă NUMAI pentru CV-uri românești fără cifre.
═══════════════════════════════════════════════════════════

Leadership cu dovadă clară ("Manager 5+ oameni, 2+ ani"): 6
Leadership implicit din responsabilități descrise (coordonare echipă, fără cifre exacte): 3
"Am spirit de lider" fără nicio dovadă: 0

Rezultate măsurabile cu cifre din CV: "Creștere 40% revenue 2024" → 5 | "Reducere costuri 15%" → 5
Rezultate descrise fără cifre (CV românesc tipic): "Coordonat proiecte de implementare", "Gestionat relații clienți cheie" → 2-3
Vag complet ("am obținut rezultate bune"): 0

Stabilitate: 5+ ani la 1-2 angajatori → 4 | Progres clar junior→senior → 3 | Job-hopper: -5
Disponibilitate declarată explicit: 2 | Vagă/nedeclarată: 0

INTERZIS acordare puncte pentru:
"Sunt comunicativ", "muncitor", "adaptabil", "orientat spre rezultate" — ZERO fără nicio dovadă.

PLAFON CV ROMÂNESC FĂRĂ KPI-URI: dacă candidatul nu are nicio cifră/rezultat cuantificat
  dar are responsabilități clare și stabilitate angajator, soft_skills_score MAX 15/20.

═══════════════════════════════════════════════════════════
FAZA 5: CALCUL SCOR FINAL (MAX 90 PUNCTE)
═══════════════════════════════════════════════════════════

SCOR = FAZA2 (0-40) + FAZA3 (0-25) + FAZA4 (0-20) + BONUS (0-5) + PENALITĂȚI (negative)
MAXIM ABSOLUT = 90 (restul de 10 puncte se câștigă la interviu față în față)

BONUS 0-5 PUNCTE (acordă NUMAI dacă există dovadă clară):
  +5: Certificări excepționale în domeniu (ex: CFA, ACCA, PMP, AWS Solutions Architect) + match exact industrie + rezultate documentate
  +3: Certificare relevantă recunoscută internațional sau experiență exactă în industria clientului
  +1-2: Formare continuă recentă demonstrată (cursuri, certificări minore relevante în ultimii 2 ani)
  0: Fără dovadă concretă de valoare adăugată excepțională

APLICARE KNOCKOUT după calcul:
  Cerință hard lipsă → final_score = MIN(calculat, 35) + recommendation = REJECT
  Supracalificat evident → final_score = MIN(calculat, 38) + overqualification_risk = HIGH
  CV total vag → final_score = MIN(calculat, 50) + must_verify_by_phone obligatoriu

CATEGORII FINALE (pe scara 0-90):
  70-90: STRONG_YES (TOP CANDIDAT) | 45-69: YES (POTENȚIAL) | sub 45: NO/REJECT (NERECOMANDAT)

ACȚIUNI RECOMANDATE:
  recommended_next_action: CALL_NOW dacă scor >= 65 | CALL_LATER dacă scor 45-64 | KEEP_FOR_OTHER_ROLE dacă < 45 dar potrivit alt rol | REJECT dacă nu are potențial real
  priority: HIGH dacă scor >= 70 | MEDIUM dacă scor 45-69 | LOW dacă < 45

═══════════════════════════════════════════════════════════
REGULI ANTI-INFLAȚIE (CRITICE — nu le ignora)
═══════════════════════════════════════════════════════════

SCOR NU CREȘTE pentru:
  ✗ CV bine formatat sau aspectuos (design ≠ competență)
  ✗ Soft skills enumerate fără dovadă ("sunt serios", "punctual")
  ✗ Presupuneri: NU acorda permis dacă nu scrie, NU acorda engleză dacă nu scrie, NU acorda IT skills dacă nu scrie

SCOR CREȘTE NUMAI pentru:
  ✓ Ani numerici cu perioadă clară: "5 ani IT, 2019-2024"
  ✓ Titlu de job + durată dovedite
  ✓ Rezultat concret cuantificat cu cifre
  ✓ Certificare/licență confirmată explicit în CV

═══════════════════════════════════════════════════════════
CÂMPURI OBLIGATORII — completează cu atenție
═══════════════════════════════════════════════════════════

candidate_name: extrage exact din CV, nu inventa
email + phone: exact din CV, sau gol dacă lipsesc
companies: lista angajatorilor menționați explicit — nu inventa
positions_held: lista funcțiilor ocupate din CV
current_position: ultimul titlu real din CV
recommended_role_for_candidate: jobul natural al candidatului bazat pe CV real — IGNORĂ postul vizat — scrie ÎNTOTDEAUNA ceva concret (ex: "Contabil", "Sofer TIR", "Vanzator", "Inginer Mecanic")
decision_reason: citează exact din CV ce justifică scorul (ex: "Are 6 ani experiență directă ca operator depozit, 2018-2024. Lipsește permisul de stivuitor menționat în cerințe.")
must_verify_by_phone: 3-5 întrebări tăioase și directe pentru clarificat ÎNAINTE de interviu
manager_summary: max 3 rânduri pentru hiring manager, fără jargon HR
phone_call_script: text gata de folosit la telefon, în română, natural, cu salut + motiv apel + 3-4 întrebări cheie + închidere
interview_questions: 4-6 întrebări specifice bazate pe CV-ul acestui candidat
reject_reason_internal: completează NUMAI dacă recommendation este NO sau REJECT
scoring_breakdown: scurt text cu calculul: "exp 36/40 + edu 10/25 + soft 8/20 - penalitati 5 = 49"

Răspunde STRICT cu JSON valid. Zero text înainte sau după JSON.

{schema}

CV CANDIDAT:
{text[:15000]}"""


def analyze_cv_with_ai(text, target_job, job_requirements=None):
    if not os.getenv("OPENROUTER_API_KEY"):
        raise ValueError("Lipseste OPENROUTER_API_KEY din .env")
    model = os.getenv("OPENROUTER_MODEL", "openai/gpt-4o-mini")
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": build_prompt(text, target_job, job_requirements)}]
    )
    return response.choices[0].message.content


def parse_ai_response(content):
    if not content:
        return None
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", content, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                return None
    return None


def clean_score(v):
    try:
        return max(0, min(90, int(v)))
    except Exception:
        return 0


def as_text(v):
    if v is None:
        return ""
    if isinstance(v, list):
        return ", ".join(str(x) for x in v)
    if isinstance(v, dict):
        return json.dumps(v, ensure_ascii=False)
    return str(v)


def filename_to_name(file):
    clean = os.path.splitext(file)[0].replace("-"," ").replace("_"," ").replace("."," ")
    bad = {"cv","c","v","curriculum","vitae","resume","formal","2024","2025","pdf"}
    words = [w for w in clean.split() if w.lower() not in bad]
    return (" ".join(words).strip() or clean).title()


def normalize_recommendation(value, score):
    value = as_text(value).upper().strip()
    mapping = {"STRONG_YES":"HIRE","YES":"HIRE","MAYBE":"CONSIDER","NO":"REJECT","REJECT":"REJECT","HIRE":"HIRE","CONSIDER":"CONSIDER"}
    if value in mapping:
        return mapping[value]
    # Scala 0-90: HIRE >= 70, CONSIDER >= 45, REJECT sub 45
    if score >= 70:
        return "HIRE"
    if score >= 45:
        return "CONSIDER"
    return "REJECT"


def recommendation_to_status(rec):
    return {"HIRE":"ADMIS","CONSIDER":"DE ANALIZAT","REJECT":"RESPINS"}.get(str(rec).upper().strip(), "NOU")


def _compute_priority(score, data):
    """Calculeaza priority local ca fallback daca AI nu returneaza corect. Scala 0-90."""
    ai_priority = as_text(data.get("priority", "")).upper().strip()
    if ai_priority in ("HIGH", "MEDIUM", "LOW"):
        return ai_priority
    if score >= 70:
        return "HIGH"
    if score >= 45:
        return "MEDIUM"
    return "LOW"


def _compute_next_action(score, data):
    """Calculeaza recommended_next_action local ca fallback. Scala 0-90."""
    ai_action = as_text(data.get("recommended_next_action", "")).upper().strip()
    if ai_action in ("CALL_NOW", "CALL_LATER", "KEEP_FOR_OTHER_ROLE", "REJECT"):
        return ai_action
    if score >= 65:
        return "CALL_NOW"
    if score >= 45:
        return "CALL_LATER"
    better = as_text(data.get("better_role_match", "")).strip()
    if better:
        return "KEEP_FOR_OTHER_ROLE"
    return "REJECT"


# ─── DOMAIN COMPATIBILITY MATRIX (100 domenii, 23 clustere, 3 niveluri) ──────
# Niveluri: 2 = COMPATIBIL | 1 = UPSKILL (max 45) | 0 = INCOMPATIBIL (max 28)

_CLUSTER_KW = [
    # ── HORECA & Turism ────────────────────────────────────────────────────────
    ("horeca", [
        "bucatar", "ospatar", "barman", "restaurant", "cofetarie", "patiserie",
        "horeca", "bucatarie", "chef", "cook", "sef bucatarie", "pizzar",
        "camerista", "housekeeping", "receptionist hotel", "turism", "ghid turistic",
        "catering", "events", "sommelier", "food service",
    ]),
    # ── Transport & Logistică ──────────────────────────────────────────────────
    ("transport", [
        "sofer", "tir", "curier", "livrari", "camion", "taximetrist", "taxi",
        "conducator auto", "livrator", "dispecer transport", "transport rutier",
        "autobuz", "microbuz", "transport persoane", "aviatie", "aeroport",
        "marina", "port", "supply chain", "forklift", "transport international",
        "logistica", "dispecer logistica",
    ]),
    # ── Depozit & Producție Entry ──────────────────────────────────────────────
    ("depozit_productie", [
        "stivuitorist", "manipulant", "picker", "packer", "lucrator depozit",
        "muncitor depozit", "operator depozit", "agent depozit", "sorter",
        "ambalator", "productie", "manufactura", "operator productie",
        "alimentara", "procesare carne", "textile", "confectii", "mobilier",
        "ambalare", "etichetare", "muncitor", "muncitor necalificat",
    ]),
    # ── Retail & Vânzări Entry ─────────────────────────────────────────────────
    ("retail_entry", [
        "casier", "vanzator", "lucrator comercial", "agent vanzari", "lucrator magazin",
        "reponitor", "operator vanzari", "merchandiser", "promotor",
        "supermarket", "hypermarket", "fashion retail", "farmacie", "droguerie",
        "vanzari b2c", "lucrator farmacie",
    ]),
    # ── Pază & Securitate ─────────────────────────────────────────────────────
    ("paza", [
        "paznic", "agent paza", "guard", "bodyguard", "supraveghetor",
        "agent securitate", "inspector paza", "securitate evenimente",
        "monitorizare", "cctv", "close protection",
    ]),
    # ── Construcții & Instalații ───────────────────────────────────────────────
    ("constructii", [
        "zidar", "faiantar", "zugrav", "tamplar", "dulgher", "instalator",
        "beton", "saper", "constructor", "montator", "lacatus", "sudor",
        "electrician constructii", "fierar beton", "muncitor constructii",
        "santier", "supraveghere santier", "topografie", "geodezie", "plumber",
    ]),
    # ── IT & Tech ─────────────────────────────────────────────────────────────
    ("it_tech", [
        "programator", "developer", "devops", "software", "it specialist",
        "web developer", "data analyst", "data scientist", "cloud", "cybersecurity",
        "network engineer", "full stack", "backend", "frontend", "system admin",
        "it manager", "data engineer", "qa engineer", "tester", "scrum master",
        "product owner", "tech lead", "mobile developer", "ai engineer",
        "machine learning", "it project manager", "it support", "helpdesk",
        "digital", "web design",
    ]),
    # ── Inginerie ─────────────────────────────────────────────────────────────
    ("inginerie", [
        "inginer", "arhitect tehnic", "proiectant", "inginer mecanic",
        "inginer electric", "inginer electronic", "inginer civil",
        "inginer industrial", "inginer auto", "automotive engineer",
        "inginer energetic", "renewables", "inginer chimica", "petrochimie",
        "inginer productie", "inginer calitate", "inginer service", "mentenanta",
        "inginer apa", "cad specialist",
    ]),
    # ── Medical & Sănătate ────────────────────────────────────────────────────
    ("medical", [
        "medic", "asistent medical", "asistenta medicala", "asistenta generalista",
        "farmacist", "infirmiera", "radiolog", "stomatolog", "kinetoterapeut",
        "fizioterapeut", "logoped", "psiholog clinic", "dentist", "moasa",
        "brancardier", "ingrijitor batrani", "home care", "laborant",
        "tehnician dentar", "psiholog", "psihoterapeut",
    ]),
    # ── Financiar & Contabilitate ─────────────────────────────────────────────
    ("financiar", [
        "contabil", "economist", "auditor", "controller", "analist financiar",
        "cfo", "trezorier", "contabilitate", "specialist fiscal", "consultant fiscal",
        "banca", "credit", "asigurari", "broker", "investitii", "chief accountant",
        "finante", "trezorerie", "analyst financiar",
    ]),
    # ── Juridic ───────────────────────────────────────────────────────────────
    ("juridic", [
        "avocat", "jurist", "consilier juridic", "notar", "judecator",
        "procuror", "paralegal", "compliance", "contract manager",
        "executor judecatoresc", "risk manager juridic", "legal",
    ]),
    # ── HR & Recrutare ────────────────────────────────────────────────────────
    ("hr_recrutare", [
        "recruiter", "specialist hr", "hr manager", "hr business partner",
        "talent acquisition", "resurse umane", "specialist recrutare",
        "hr generalist", "payroll", "admin hr", "trainer corporativ",
        "learning development", "l&d", "hr specialist",
    ]),
    # ── Educație & Training ───────────────────────────────────────────────────
    ("educatie", [
        "profesor", "invatator", "educator", "instructor", "formator",
        "trainer", "lector", "mentor", "coach", "educator gradinita",
        "profesor universitar", "cadru didactic",
    ]),
    # ── Administrativ / Secretariat ───────────────────────────────────────────
    ("administrativ", [
        "secretara", "asistent administrativ", "receptionist", "operator date",
        "coordonator administrativ", "office manager", "asistent manager",
        "back office", "administrativ",
    ]),
    # ── Call Center / BPO / Customer Support ──────────────────────────────────
    ("call_center", [
        "call center", "customer support", "customer service", "bpo",
        "agent suport", "operator call center", "helpdesk client",
        "relatii clienti", "agent relatii",
    ]),
    # ── Marketing / Digital Marketing ─────────────────────────────────────────
    ("marketing", [
        "marketing", "digital marketing", "social media", "seo", "sem",
        "content creator", "copywriter", "brand manager", "pr", "comunicare",
        "publicitate", "media", "grafician", "ux", "ui designer",
    ]),
    # ── Sales B2B & Account Management ────────────────────────────────────────
    ("sales_b2b", [
        "account manager", "sales manager", "b2b sales", "territory manager",
        "key account", "sales representative", "business developer",
        "sales engineer", "vanzari b2b", "reprezentant comercial",
    ]),
    # ── Imobiliare ────────────────────────────────────────────────────────────
    ("imobiliare", [
        "imobiliare", "agent imobiliar", "broker imobiliar", "evaluator imobiliar",
        "property manager", "real estate",
    ]),
    # ── Beauty & Wellness ─────────────────────────────────────────────────────
    ("beauty_wellness", [
        "cosmeticiana", "frizer", "coafor", "manichiura", "pedichiura",
        "make-up", "estetician", "spa", "maseur", "masaj", "beauty",
        "nail technician", "wellness",
    ]),
    # ── Agricultură ───────────────────────────────────────────────────────────
    ("agricultura", [
        "agricultor", "fermier", "zootehnist", "agronom", "viticultor",
        "horticultor", "muncitor agricol", "operator utilaje agricole", "silvicultura",
    ]),
    # ── Automotive Service ────────────────────────────────────────────────────
    ("automotive", [
        "mecanic auto", "tinichigiu", "vopsitor auto", "electician auto",
        "diagnostician auto", "tehnician auto", "service auto", "vulcanizare",
        "inspector itp", "mecanic utilaje",
    ]),
    # ── Management Executive ──────────────────────────────────────────────────
    ("management", [
        "manager general", "director general", "ceo", "coo", "cto", "cfo exec",
        "director executiv", "director operational", "vp", "vice president",
        "country manager", "general manager",
    ]),
    # ── ONG / Public / Administratie Publica ──────────────────────────────────
    ("ong_public", [
        "ong", "ngo", "asociatie", "fundatie", "functionar public",
        "administratie publica", "primarie", "prefectura", "minister",
        "institutie publica", "sector public", "referent", "inspector public",
    ]),
]

# ── 7 Super-clustere ──────────────────────────────────────────────────────────
_SUPER = {
    "horeca":           "blue_collar",
    "transport":        "blue_collar",
    "depozit_productie":"blue_collar",
    "retail_entry":     "blue_collar",
    "paza":             "blue_collar",
    "agricultura":      "blue_collar",
    "constructii":      "constructii_tehnic",
    "automotive":       "constructii_tehnic",
    "it_tech":          "tech_ing",
    "inginerie":        "tech_ing",
    "medical":          "medical",
    "beauty_wellness":  "medical",
    "financiar":        "business",
    "juridic":          "business",
    "hr_recrutare":     "business",
    "management":       "business",
    "imobiliare":       "business",
    "sales_b2b":        "business",
    "call_center":      "support",
    "administrativ":    "support",
    "marketing":        "support",
    "ong_public":       "support",
    "educatie":         "educatie",
}

#             blue  cstr  tech  med   biz   supp  edu
_SC_COMPAT = {
    ("blue_collar",      "blue_collar"):       2,
    ("blue_collar",      "constructii_tehnic"):1,
    ("blue_collar",      "tech_ing"):          0,
    ("blue_collar",      "medical"):           0,
    ("blue_collar",      "business"):          0,
    ("blue_collar",      "support"):           1,
    ("blue_collar",      "educatie"):          0,

    ("constructii_tehnic","blue_collar"):      1,
    ("constructii_tehnic","constructii_tehnic"):2,
    ("constructii_tehnic","tech_ing"):         1,
    ("constructii_tehnic","medical"):          0,
    ("constructii_tehnic","business"):         0,
    ("constructii_tehnic","support"):          1,
    ("constructii_tehnic","educatie"):         0,

    ("tech_ing",         "blue_collar"):       0,
    ("tech_ing",         "constructii_tehnic"):1,
    ("tech_ing",         "tech_ing"):          2,
    ("tech_ing",         "medical"):           0,
    ("tech_ing",         "business"):          2,
    ("tech_ing",         "support"):           1,
    ("tech_ing",         "educatie"):          1,

    ("medical",          "blue_collar"):       0,
    ("medical",          "constructii_tehnic"):0,
    ("medical",          "tech_ing"):          0,
    ("medical",          "medical"):           2,
    ("medical",          "business"):          0,
    ("medical",          "support"):           0,
    ("medical",          "educatie"):          1,

    ("business",         "blue_collar"):       0,
    ("business",         "constructii_tehnic"):0,
    ("business",         "tech_ing"):          2,
    ("business",         "medical"):           0,
    ("business",         "business"):          2,
    ("business",         "support"):           2,
    ("business",         "educatie"):          2,

    ("support",          "blue_collar"):       1,
    ("support",          "constructii_tehnic"):1,
    ("support",          "tech_ing"):          1,
    ("support",          "medical"):           0,
    ("support",          "business"):          2,
    ("support",          "support"):           2,
    ("support",          "educatie"):          2,

    ("educatie",         "blue_collar"):       0,
    ("educatie",         "constructii_tehnic"):0,
    ("educatie",         "tech_ing"):          1,
    ("educatie",         "medical"):           1,
    ("educatie",         "business"):          2,
    ("educatie",         "support"):           2,
    ("educatie",         "educatie"):          2,
}


def _detect_cluster(text):
    """Detecteaza cluster-ul profesional dintr-un titlu de job."""
    t = normalize_text(text)
    for cluster, kws in _CLUSTER_KW:
        for kw in kws:
            if kw in t:
                return cluster
    return None


def _get_compat_level(natural_role, target_job, cv_text="", candidate_cluster=None, job_cluster=None):
    """
    Returneaza nivelul de compatibilitate domeniu:
      2 = COMPATIBIL (fara limitare)
      1 = UPSKILL (scor max 45)
      0 = INCOMPATIBIL (scor max 28 + REJECT)

    Prioritate 1: clustere clasificate de AI (candidate_cluster / job_cluster din DB)
    Prioritate 2: keyword detection ca fallback pentru inregistrari vechi
    """
    ns, ts = None, None

    # Prioritate 1: clustere AI din DB (mai precise decat keyword matching)
    if candidate_cluster and job_cluster:
        cc = as_text(candidate_cluster).lower().strip()
        jc = as_text(job_cluster).lower().strip()
        if cc in VALID_CLUSTERS and jc in VALID_CLUSTERS:
            ns, ts = cc, jc

    # Prioritate 2: keyword detection fallback
    if not ns or not ts:
        nc = _detect_cluster(natural_role)
        tc = _detect_cluster(target_job)
        if not nc or not tc:
            return 2  # domeniu necunoscut = nu penaliza
        if nc == tc:
            return 2  # acelasi cluster = compatibil
        ns = _SUPER.get(nc, "")
        ts = _SUPER.get(tc, "")
        if not ns or not ts:
            return 2

    if ns == ts:
        return 2  # acelasi super-cluster = compatibil

    level = _SC_COMPAT.get((ns, ts), 2)

    # Exceptie: constructii_tehnic → tech_ing cu studii superioare in CV
    if ns == "constructii_tehnic" and ts == "tech_ing" and level == 1:
        cv_n = normalize_text(cv_text or "")
        edu_kw = ["licenta", "master", "facultate", "universitate", "inginer", "tehnica constructii", "politehn", "utcb", "utcn"]
        if any(kw in cv_n for kw in edu_kw):
            return 2  # studii detectate = compatibil deplin

    return level


def _is_domain_incompatible(natural_role, target_job, cv_text=""):
    """Backward compat: returneaza True daca nivel 0 (incompatibil)."""
    return _get_compat_level(natural_role, target_job, cv_text) == 0


VALID_CLUSTERS = {"blue_collar", "constructii_tehnic", "tech_ing", "medical", "business", "support", "educatie"}


def _validate_cluster(val, fallback_text=""):
    """Valideaza si returneaza un super-cluster valid, cu fallback la keyword detection."""
    v = as_text(val).lower().strip().replace("-", "_").replace(" ", "_")
    if v in VALID_CLUSTERS:
        return v
    cl = _detect_cluster(fallback_text)
    if cl:
        sc = _SUPER.get(cl)
        if sc in VALID_CLUSTERS:
            return sc
    return None


def apply_local_safety_rules(data, target_job, cv_text, profile, candidate_cluster=None, job_cluster=None):
    cv = normalize_text(cv_text)
    combined = normalize_text(" ".join([as_text(data.get(k,"")) for k in ["name","position","summary","strengths","skills"]]) + " " + cv[:3000])

    # ── Regula 1: supracalificare academica pentru rol operational ─────────────
    operational_domains = ["Depozit Logistica","Retail Financiar","Retail","Transport Soferie","Transport Livrari","Productie","Constructii","HoReCa","Facility","Securitate"]
    high_level = ["rector","profesor universitar","director general","ceo","antreprenor","consultant strategic","decan","academic"]
    if profile.get("domain") in operational_domains and any(w in combined for w in high_level):
        data["score"] = min(clean_score(data.get("score",0)), 35)
        data["recommendation"] = "REJECT"
        data["summary"] = set_summary_risk(data.get("summary",""), over_risk="high", level_risk="high")
        data["priority"] = "LOW"
        data["recommended_next_action"] = "REJECT"
        data["reject_reason_internal"] = "Supracalificare ridicata pentru rol operational."

    # ── Regula 2: compatibilitate domeniu (3 niveluri) ────────────────────────
    natural_role = as_text(data.get("position", ""))
    if natural_role:
        compat = _get_compat_level(natural_role, target_job, cv_text,
                                   candidate_cluster=candidate_cluster,
                                   job_cluster=job_cluster)
        nc = _detect_cluster(natural_role)
        tc = _detect_cluster(target_job)

        if compat == 0:  # INCOMPATIBIL → max 28 + REJECT
            if clean_score(data.get("score", 0)) > 28:
                data["score"] = 28
            data["recommendation"] = "REJECT"
            data["priority"] = "LOW"
            data["recommended_next_action"] = "REJECT"
            data["level_mismatch"] = "HIGH"
            data["reject_reason_internal"] = (
                f"Incompatibilitate domeniu: '{nc}' vs post '{target_job}' ({tc}). "
                "Transfer de cariera fara dovezi de reconversie."
            )
            clean_visible = strip_risk_text(data.get("summary", ""))
            data["summary"] = set_summary_risk(
                f"Domeniu incompatibil ({nc} vs {tc}). {clean_visible}",
                over_risk="low", level_risk="high"
            )

        elif compat == 1:  # UPSKILL → max 45 + pastram ca CONSIDER
            if clean_score(data.get("score", 0)) > 45:
                data["score"] = 45
            if data.get("recommendation") not in ("CONSIDER", "REJECT"):
                data["recommendation"] = "CONSIDER"
            if data.get("priority") == "HIGH":
                data["priority"] = "MEDIUM"
            if data.get("recommended_next_action") == "CALL_NOW":
                data["recommended_next_action"] = "CALL_LATER"
            data["level_mismatch"] = "MEDIUM"
    req = detect_required_license(target_job, profile)
    if req:
        r = req.lower()
        found = (f"permis {r}" in cv or f"categoria {r}" in cv or f"cat {r}" in cv or (req=="CE" and "tir" in cv))
        if not found:
            data["score"] = min(clean_score(data.get("score",0)), 30)
            data["recommendation"] = "REJECT"
            clean_visible = strip_risk_text(data.get("summary",""))
            data["summary"] = set_summary_risk(f"Lipseste confirmarea clara pentru permis categoria {req}. {clean_visible}", over_risk="high", level_risk="high")
            data["priority"] = "LOW"
            data["recommended_next_action"] = "REJECT"
            missing = as_text(data.get("missing_requirements", ""))
            permis_item = f"Nu apare permis categoria {req} in CV"
            if permis_item not in missing:
                existing = [x.strip() for x in missing.split(",") if x.strip()] if missing else []
                existing.insert(0, permis_item)
                data["missing_requirements"] = existing
    return data


def strip_risk_text(value):
    text = as_text(value).strip()
    text = re.sub(r"\bRisc\s+supracalificare\s*:\s*(low|medium|high|unknown|scazut|mediu|ridicat|necunoscut)\s*\.?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bNepotrivire\s+nivel\s*:\s*(low|medium|high|unknown|scazut|mediu|ridicat|necunoscut)\s*\.?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\boverqualification_risk\s*:\s*(low|medium|high|unknown)\s*\.?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\blevel_mismatch\s*:\s*(low|medium|high|unknown)\s*\.?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def set_summary_risk(summary, over_risk=None, level_risk=None):
    text = as_text(summary)
    over = over_risk or "unknown"
    level = level_risk or "unknown"

    current_over = re.search(r"Risc supracalificare:\s*(low|medium|high|unknown)", text, flags=re.IGNORECASE)
    current_level = re.search(r"Nepotrivire nivel:\s*(low|medium|high|unknown)", text, flags=re.IGNORECASE)

    if current_over:
        text = re.sub(r"Risc supracalificare:\s*(low|medium|high|unknown)", f"Risc supracalificare: {over}", text, flags=re.IGNORECASE)
    else:
        text = f"Risc supracalificare: {over}. " + text

    if current_level:
        text = re.sub(r"Nepotrivire nivel:\s*(low|medium|high|unknown)", f"Nepotrivire nivel: {level}", text, flags=re.IGNORECASE)
    else:
        parts = text.split(".", 1)
        if len(parts) == 2 and "Risc supracalificare:" in parts[0]:
            text = parts[0] + f". Nepotrivire nivel: {level}." + parts[1]
        else:
            text = f"Nepotrivire nivel: {level}. " + text

    return re.sub(r"\s+", " ", text).strip()


def first_value(*values):
    for value in values:
        text = as_text(value).strip()
        if text:
            return text
    return ""


def normalize_data(data, file, target_job, cv_text, profile):
    score = clean_score(data.get("final_score", data.get("fit_percentage", 0)))
    rec = normalize_recommendation(data.get("recommendation",""), score)
    name = data.get("candidate_name") or data.get("name") or data.get("full_name") or filename_to_name(file)

    over_risk = as_text(data.get("overqualification_risk", "unknown")).lower() or "unknown"
    level_risk = as_text(data.get("level_mismatch", "unknown")).lower() or "unknown"
    instability_risk = as_text(data.get("instability_risk", "unknown")).lower() or "unknown"
    cv_quality = as_text(data.get("cv_quality_score", "")).upper() or "DECENT"

    clean_ai_summary = strip_risk_text(data.get("summary", data.get("match_analysis", "")))
    if not clean_ai_summary:
        clean_ai_summary = "Rezumat indisponibil. Verifica manual CV-ul."

    summary = f"Risc supracalificare: {over_risk}. Nepotrivire nivel: {level_risk}. Stabilitate: {instability_risk}. CV: {cv_quality}. {clean_ai_summary}"

    recommended_role = first_value(
        data.get("recommended_role_for_candidate"),
        data.get("current_position"),
        data.get("target_position"),
        data.get("position")
    )

    priority = _compute_priority(score, data)
    recommended_next_action = _compute_next_action(score, data)

    # scoring_breakdown pentru transparenta
    scoring_bd = as_text(data.get("scoring_breakdown", ""))
    if not scoring_bd:
        exp = data.get("experience_score", 0)
        edu = data.get("education_score", 0)
        soft = data.get("soft_skills_score", 0)
        pen = data.get("penalties_total", 0)
        bon = data.get("bonuses_total", 0)
        if any([exp, edu, soft, pen, bon]):
            scoring_bd = f"exp {exp}/40 + edu {edu}/25 + soft {soft}/20 + bonus {bon}/5 - pen {pen} = {score}"

    # Clustere domeniu clasificate de AI
    candidate_cluster = _validate_cluster(
        data.get("candidate_cluster", ""),
        fallback_text=recommended_role
    )
    job_cluster = _validate_cluster(
        data.get("job_cluster", ""),
        fallback_text=target_job
    )

    normalized = {
        "name": name,
        "email": data.get("email",""),
        "phone": data.get("phone",""),
        "position": recommended_role,
        "years_experience": data.get("years_experience",""),
        "skills": data.get("transferable_skills", data.get("skills","")),
        "companies": data.get("companies",""),
        "score": score,
        "level": data.get("confidence_level", data.get("level","")),
        "strengths": data.get("strengths",""),
        "weaknesses": data.get("gaps", data.get("weaknesses","")),
        "summary": summary,
        "recommendation": rec,
        "decision_reason":         as_text(data.get("decision_reason", "")),
        "missing_requirements":    as_text(data.get("missing_requirements", [])),
        "must_verify_by_phone":    as_text(data.get("must_verify_by_phone", [])),
        "recommended_next_action": recommended_next_action,
        "priority":                priority,
        "manager_summary":         as_text(data.get("manager_summary", "")),
        "phone_call_script":       as_text(data.get("phone_call_script", "")),
        "interview_questions":     as_text(data.get("interview_questions", [])),
        "better_role_match":       as_text(data.get("better_role_match", "")),
        "reject_reason_internal":  as_text(data.get("reject_reason_internal", "")),
        "candidate_cluster":       candidate_cluster,
        "job_cluster":             job_cluster,
    }
    return apply_local_safety_rules(normalized, target_job, cv_text, profile,
                                    candidate_cluster=candidate_cluster,
                                    job_cluster=job_cluster)


def save_candidate_to_db(data, file, target_job):
    try:
        db = SessionLocal()
        status = recommendation_to_status(data.get("recommendation", "CONSIDER"))
        candidate = Candidate(
            name=as_text(data.get("name","")),
            email=as_text(data.get("email","")),
            phone=as_text(data.get("phone","")),
            position=as_text(data.get("position","")),
            experience=as_text(data.get("years_experience","")),
            skills=as_text(data.get("skills","")),
            companies=as_text(data.get("companies","")),
            score=clean_score(data.get("score",0)),
            level=as_text(data.get("level","")),
            strengths=as_text(data.get("strengths","")),
            weaknesses=as_text(data.get("weaknesses","")),
            summary=as_text(data.get("summary","")),
            job_title=target_job,
            status=status,
            cv_file=file,
            # ── Campuri noi ──
            decision_reason=as_text(data.get("decision_reason","")),
            missing_requirements=as_text(data.get("missing_requirements","")),
            must_verify_by_phone=as_text(data.get("must_verify_by_phone","")),
            recommended_next_action=as_text(data.get("recommended_next_action","")),
            priority=as_text(data.get("priority","")),
            manager_summary=as_text(data.get("manager_summary","")),
            phone_call_script=as_text(data.get("phone_call_script","")),
            interview_questions=as_text(data.get("interview_questions","")),
            better_role_match=as_text(data.get("better_role_match","")),
            reject_reason_internal=as_text(data.get("reject_reason_internal","")),
            candidate_cluster=data.get("candidate_cluster") or None,
            job_cluster=data.get("job_cluster") or None,
        )
        db.add(candidate)
        db.commit()
        candidate_id = candidate.id
        db.close()
        print("✓ SALVAT IN BAZA DE DATE:", data.get("name",""), "-", status)
        return candidate_id
    except Exception as e:
        print("EROARE LA SALVARE IN DB:", e)
        try:
            db.rollback()
            db.close()
        except Exception:
            pass
        return None


# ─── PARALLEL TASK ────────────────────────────────────────────────────────────

def _process_single_cv(args):
    file, target_job, profile, job_requirements = args

    cached = get_cached_analysis(file, target_job)
    if cached:
        logger.info(f"CACHE HIT: {file} pentru {target_job}")
        return {"file": file, "data": cached, "from_cache": True, "error": None}

    path = os.path.join(UPLOAD_FOLDER, file)
    if not os.path.exists(path):
        return {"file": file, "data": None, "from_cache": False, "error": "Fisier negasit local"}

    text = extract_text_from_file(path)
    if not text:
        return {"file": file, "data": None, "from_cache": False, "error": "Nu s-a putut extrage text"}

    try:
        ai_response = analyze_cv_with_ai(text, target_job, job_requirements)
        data = parse_ai_response(ai_response)
        if not data:
            return {"file": file, "data": None, "from_cache": False, "error": "Raspuns AI invalid"}

        normalized = normalize_data(data, file, target_job, text, profile)
        save_analysis_to_cache(file, target_job, normalized)

        return {"file": file, "data": normalized, "from_cache": False, "error": None}

    except Exception as e:
        return {"file": file, "data": None, "from_cache": False, "error": str(e)}


def _get_already_analyzed_files(target_job):
    """Returneaza setul de cv_file-uri deja salvate in Candidate DB pentru acest job."""
    db = SessionLocal()
    try:
        rows = db.query(Candidate.cv_file).filter(
            Candidate.job_title == target_job,
            Candidate.cv_file != None,
            Candidate.cv_file != ""
        ).all()
        return {r.cv_file for r in rows}
    except Exception as e:
        logger.warning(f"Nu s-au putut verifica candidatii existenti: {e}")
        return set()
    finally:
        db.close()


# ─── MAIN ENTRY POINT ─────────────────────────────────────────────────────────

def process_cvs_for_job(target_job, job_requirements=None):
    target_job = target_job.strip()
    if not target_job:
        return {"ok": False, "message": "Nu ai completat postul.", "saved": 0}

    profile = match_job_profile(target_job) or build_fallback_profile(target_job)
    print("Profil job folosit:", profile.get("job_title"), "-", profile.get("domain"))
    if job_requirements:
        print("Cerinte post din platforma: DA")

    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    synced = sync_from_cloudinary()
    if synced > 0:
        print(f"Sincronizare Cloudinary: {synced} fisiere descarcate local")

    try:
        all_files = [
            f for f in os.listdir(UPLOAD_FOLDER)
            if f.lower().endswith((".pdf", ".docx", ".doc"))
        ]
    except FileNotFoundError:
        all_files = []

    if not all_files:
        return {"ok": False, "message": "Nu exista CV-uri. Verifica Cloudinary sau incarca CV-uri.", "saved": 0}

    # Pre-check: care CV-uri sunt deja analizate si salvate in DB pentru acest job?
    already_in_db = _get_already_analyzed_files(target_job)
    new_files = [f for f in all_files if f not in already_in_db]
    db_skipped = len(already_in_db.intersection(set(all_files)))

    print(f"\nTotal CV-uri: {len(all_files)} | Deja in DB: {db_skipped} | De analizat: {len(new_files)}")
    print("=" * 70)

    cache_hits = 0
    fresh_analyses = 0
    saved = 0
    results = []

    if new_files:
        tasks = [(file, target_job, profile, job_requirements) for file in new_files]

        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = {executor.submit(_process_single_cv, task): task[0] for task in tasks}

            for future in as_completed(futures):
                result = future.result()
                file = result["file"]

                if result["error"]:
                    print(f"✗ EROARE {file}: {result['error']}")
                    continue

                data = result["data"]
                if not data:
                    continue

                if result["from_cache"]:
                    cache_hits += 1
                    print(f"⚡ CACHE: {file} -> {data.get('name')} (scor: {data.get('score')})")
                else:
                    fresh_analyses += 1
                    print(f"✓ ANALIZAT: {file} -> {data.get('name')} (scor: {data.get('score')})")

                save_candidate_to_db(data, file, target_job)
                saved += 1
                results.append({
                    "name": data.get("name"),
                    "score": data.get("score"),
                    "recommendation": data.get("recommendation"),
                    "priority": data.get("priority"),
                    "recommended_next_action": data.get("recommended_next_action"),
                })

    results.sort(key=lambda x: x.get("score", 0), reverse=True)

    print("\n" + "=" * 70)
    print("SUMMARY:")
    print(f"  Deja in baza de date: {db_skipped}")
    print(f"  Din cache (instant):  {cache_hits}")
    print(f"  Analizate acum (AI):  {fresh_analyses}")
    print(f"  Noi salvati:          {saved}")
    print("=" * 70)

    parts = []
    if fresh_analyses > 0:
        parts.append(f"{fresh_analyses} CV-uri noi analizate")
    if cache_hits > 0:
        parts.append(f"{cache_hits} preluate din cache")
    if db_skipped > 0:
        parts.append(f"{db_skipped} deja existente in baza de date")

    total_processed = db_skipped + saved
    message = f"Analiza finalizata. {', '.join(parts)}. Total candidati pentru acest post: {total_processed}."

    return {
        "ok": True,
        "message": message,
        "saved": saved,
        "cache_hits": cache_hits,
        "fresh_analyses": fresh_analyses,
        "results": results
    }


def parse_all_cvs():
    target_job = input("Pentru ce post cauti candidat? ")
    process_cvs_for_job(target_job)


if __name__ == "__main__":
    parse_all_cvs()
